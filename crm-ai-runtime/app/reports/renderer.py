import html
import io
import re
from dataclasses import dataclass, field
from datetime import datetime
from typing import Any


@dataclass
class ReportBlock:
    kind: str
    text: str = ""
    level: int = 0
    rows: list[list[str]] = field(default_factory=list)
    items: list[str] = field(default_factory=list)
    ordered: bool = False


class MarkdownReportParser:
    _heading = re.compile(r"^(#{1,6})\s+(.+)$")
    _unordered = re.compile(r"^\s*[-*+]\s+(.+)$")
    _ordered = re.compile(r"^\s*\d+[.)]\s+(.+)$")
    _table_separator = re.compile(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$")

    def parse(self, value: str) -> list[ReportBlock]:
        lines = str(value or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
        blocks: list[ReportBlock] = []
        index = 0
        while index < len(lines):
            line = lines[index].rstrip()
            stripped = line.strip()
            if not stripped:
                index += 1
                continue
            if stripped.startswith("```"):
                language = stripped[3:].strip()
                index += 1
                values = []
                while index < len(lines) and not lines[index].strip().startswith("```"):
                    values.append(lines[index])
                    index += 1
                if index < len(lines):
                    index += 1
                blocks.append(ReportBlock(kind="code", text="\n".join(values), level=1 if language else 0))
                continue
            heading = self._heading.match(stripped)
            if heading:
                blocks.append(ReportBlock(kind="heading", text=heading.group(2).strip(), level=len(heading.group(1))))
                index += 1
                continue
            if self._is_rule(stripped):
                blocks.append(ReportBlock(kind="rule"))
                index += 1
                continue
            if index + 1 < len(lines) and "|" in stripped and self._table_separator.match(lines[index + 1].strip()):
                rows = [self._table_cells(stripped)]
                index += 2
                while index < len(lines) and "|" in lines[index] and lines[index].strip():
                    rows.append(self._table_cells(lines[index].strip()))
                    index += 1
                blocks.append(ReportBlock(kind="table", rows=rows))
                continue
            unordered = self._unordered.match(line)
            ordered = self._ordered.match(line)
            if unordered or ordered:
                is_ordered = ordered is not None
                items = []
                while index < len(lines):
                    current = self._ordered.match(lines[index]) if is_ordered else self._unordered.match(lines[index])
                    if current is None:
                        break
                    items.append(current.group(1).strip())
                    index += 1
                blocks.append(ReportBlock(kind="list", items=items, ordered=is_ordered))
                continue
            if stripped.startswith(">"):
                values = []
                while index < len(lines) and lines[index].strip().startswith(">"):
                    values.append(lines[index].strip()[1:].strip())
                    index += 1
                blocks.append(ReportBlock(kind="quote", text=" ".join(values)))
                continue
            values = [stripped]
            index += 1
            while index < len(lines) and lines[index].strip() and not self._starts_block(lines, index):
                values.append(lines[index].strip())
                index += 1
            blocks.append(ReportBlock(kind="paragraph", text=" ".join(values)))
        return blocks

    def _starts_block(self, lines: list[str], index: int) -> bool:
        value = lines[index].strip()
        if not value:
            return True
        if value.startswith("```") or value.startswith(">"):
            return True
        if self._heading.match(value) or self._unordered.match(value) or self._ordered.match(value):
            return True
        if self._is_rule(value):
            return True
        return index + 1 < len(lines) and "|" in value and self._table_separator.match(lines[index + 1].strip()) is not None

    def _table_cells(self, value: str) -> list[str]:
        text = value.strip().strip("|")
        return [item.strip() for item in text.split("|")]

    def _is_rule(self, value: str) -> bool:
        return bool(re.fullmatch(r"\s*([-*_])(?:\s*\1){2,}\s*", value))


class ReportRenderer:
    def __init__(self):
        self.parser = MarkdownReportParser()

    def render(self, title: str, content: str, report_format: str) -> tuple[bytes, str, str]:
        normalized = str(report_format or "").strip().lower()
        blocks = self.parser.parse(content)
        if normalized == "html":
            return self._render_html(title, blocks), "text/html; charset=utf-8", ".html"
        if normalized in {"docx", "word"}:
            return self._render_docx(title, blocks), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", ".docx"
        if normalized == "pdf":
            return self._render_pdf(title, blocks), "application/pdf", ".pdf"
        raise ValueError("报告格式仅支持docx、pdf和html")

    def _render_html(self, title: str, blocks: list[ReportBlock]) -> bytes:
        body = []
        for block in blocks:
            if block.kind == "heading":
                level = min(max(block.level, 1), 6)
                body.append("<h%s>%s</h%s>" % (level, self._inline_html(block.text), level))
            elif block.kind == "paragraph":
                body.append("<p>%s</p>" % self._inline_html(block.text))
            elif block.kind == "quote":
                body.append("<blockquote>%s</blockquote>" % self._inline_html(block.text))
            elif block.kind == "rule":
                body.append("<hr>")
            elif block.kind == "code":
                body.append("<pre><code>%s</code></pre>" % html.escape(block.text))
            elif block.kind == "list":
                tag = "ol" if block.ordered else "ul"
                items = "".join("<li>%s</li>" % self._inline_html(item) for item in block.items)
                body.append("<%s>%s</%s>" % (tag, items, tag))
            elif block.kind == "table" and block.rows:
                head = "".join("<th>%s</th>" % self._inline_html(item) for item in block.rows[0])
                rows = []
                for row in block.rows[1:]:
                    rows.append("<tr>%s</tr>" % "".join("<td>%s</td>" % self._inline_html(item) for item in row))
                body.append("<div class=\"table-wrap\"><table><thead><tr>%s</tr></thead><tbody>%s</tbody></table></div>" % (head, "".join(rows)))
        document = """<!doctype html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{title}</title>
<style>
@page {{ size: A4; margin: 20mm 18mm; }}
* {{ box-sizing: border-box; }}
body {{ max-width: 920px; margin: 0 auto; padding: 38px 48px 60px; color: #243044; background: #fff; font: 15px/1.75 "Microsoft YaHei", "PingFang SC", sans-serif; }}
.report-head {{ margin-bottom: 30px; padding-bottom: 20px; border-bottom: 2px solid #e7ebf2; }}
.report-head h1 {{ margin: 0 0 8px; color: #152033; font-size: 30px; line-height: 1.3; }}
.report-head p {{ margin: 0; color: #7a8496; font-size: 13px; }}
h1, h2, h3, h4 {{ color: #17233a; line-height: 1.45; break-after: avoid; }}
h1 {{ margin-top: 32px; font-size: 26px; }} h2 {{ margin-top: 28px; font-size: 21px; }} h3 {{ margin-top: 24px; font-size: 17px; }}
p {{ margin: 10px 0; }} ul, ol {{ padding-left: 25px; }} li {{ margin: 5px 0; }}
blockquote {{ margin: 16px 0; padding: 11px 16px; border-left: 4px solid #5b7cfa; border-radius: 0 8px 8px 0; background: #f5f7fb; color: #4f5d73; }}
.table-wrap {{ margin: 16px 0; overflow-x: auto; }} table {{ width: 100%; border-collapse: collapse; font-size: 14px; }}
th, td {{ padding: 10px 12px; border: 1px solid #dfe4ec; text-align: left; vertical-align: top; }} th {{ background: #f1f4f9; color: #25324a; }}
pre {{ overflow-x: auto; padding: 14px 16px; border-radius: 8px; background: #172033; color: #eef2ff; font: 13px/1.6 Consolas, monospace; }}
code {{ padding: 2px 5px; border-radius: 4px; background: #f1f4f9; }} pre code {{ padding: 0; background: transparent; }}
a {{ color: #315dcc; }} hr {{ margin: 26px 0; border: 0; border-top: 1px solid #dfe4ec; }}
@media print {{ body {{ max-width: none; padding: 0; }} }}
</style>
</head>
<body>
<header class="report-head"><h1>{title}</h1><p>生成时间：{generated_at}</p></header>
<main>{body}</main>
</body>
</html>""".format(
            title=html.escape(title),
            generated_at=datetime.now().strftime("%Y-%m-%d %H:%M"),
            body="\n".join(body),
        )
        return document.encode("utf-8")

    def _render_docx(self, title: str, blocks: list[ReportBlock]) -> bytes:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt, RGBColor

        document = Document()
        section = document.sections[0]
        section.top_margin = Cm(1.9)
        section.bottom_margin = Cm(1.9)
        section.left_margin = Cm(2.1)
        section.right_margin = Cm(2.1)
        normal = document.styles["Normal"]
        normal.font.name = "Microsoft YaHei"
        normal.font.size = Pt(10.5)
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.name = "Microsoft YaHei"
        title_run.font.size = Pt(22)
        title_run.font.color.rgb = RGBColor(23, 35, 58)
        title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        generated = document.add_paragraph()
        generated.alignment = WD_ALIGN_PARAGRAPH.CENTER
        generated_run = generated.add_run("生成时间：" + datetime.now().strftime("%Y-%m-%d %H:%M"))
        generated_run.font.size = Pt(9)
        generated_run.font.color.rgb = RGBColor(122, 132, 150)

        for block in blocks:
            if block.kind == "heading":
                level = min(max(block.level, 1), 4)
                paragraph = document.add_heading(level=level)
                self._add_docx_inline(paragraph, block.text)
            elif block.kind == "paragraph":
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(6)
                self._add_docx_inline(paragraph, block.text)
            elif block.kind == "quote":
                paragraph = document.add_paragraph(style="Intense Quote")
                self._add_docx_inline(paragraph, block.text)
            elif block.kind == "rule":
                paragraph = document.add_paragraph()
                border = OxmlElement("w:pBdr")
                bottom = OxmlElement("w:bottom")
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "6")
                bottom.set(qn("w:color"), "DDE3EC")
                border.append(bottom)
                paragraph._p.get_or_add_pPr().append(border)
            elif block.kind == "code":
                paragraph = document.add_paragraph()
                run = paragraph.add_run(block.text)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            elif block.kind == "list":
                style = "List Number" if block.ordered else "List Bullet"
                for item in block.items:
                    paragraph = document.add_paragraph(style=style)
                    self._add_docx_inline(paragraph, item)
            elif block.kind == "table" and block.rows:
                column_count = max(len(row) for row in block.rows)
                table = document.add_table(rows=len(block.rows), cols=column_count)
                table.style = "Table Grid"
                for row_index, row in enumerate(block.rows):
                    for column_index in range(column_count):
                        value = row[column_index] if column_index < len(row) else ""
                        cell = table.cell(row_index, column_index)
                        cell.text = ""
                        self._add_docx_inline(cell.paragraphs[0], value, bold=row_index == 0)
                        if row_index == 0:
                            shading = OxmlElement("w:shd")
                            shading.set(qn("w:fill"), "EEF2F7")
                            cell._tc.get_or_add_tcPr().append(shading)
                document.add_paragraph()

        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run("AI 生成报告 · 请结合真实业务情况复核")
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(140, 148, 162)
        output = io.BytesIO()
        document.save(output)
        return output.getvalue()

    def _render_pdf(self, title: str, blocks: list[ReportBlock]) -> bytes:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.platypus import (
            HRFlowable,
            ListFlowable,
            ListItem,
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )

        font_name = "STSong-Light"
        try:
            pdfmetrics.getFont(font_name)
        except KeyError:
            pdfmetrics.registerFont(UnicodeCIDFont(font_name))
        pdfmetrics.registerFontFamily(
            font_name,
            normal=font_name,
            bold=font_name,
            italic=font_name,
            boldItalic=font_name,
        )
        output = io.BytesIO()
        document = SimpleDocTemplate(
            output,
            pagesize=A4,
            rightMargin=18 * mm,
            leftMargin=18 * mm,
            topMargin=18 * mm,
            bottomMargin=18 * mm,
            title=title,
            author="CRM AI Runtime",
        )
        styles = getSampleStyleSheet()
        normal = ParagraphStyle("ReportNormal", parent=styles["Normal"], fontName=font_name, fontSize=10, leading=17, textColor=colors.HexColor("#243044"), spaceAfter=6)
        title_style = ParagraphStyle("ReportTitle", parent=normal, fontSize=22, leading=30, alignment=TA_CENTER, textColor=colors.HexColor("#17233A"), spaceAfter=8)
        meta_style = ParagraphStyle("ReportMeta", parent=normal, fontSize=8.5, leading=13, alignment=TA_CENTER, textColor=colors.HexColor("#7A8496"), spaceAfter=18)
        heading_styles = {
            1: ParagraphStyle("ReportH1", parent=normal, fontSize=17, leading=24, textColor=colors.HexColor("#17233A"), spaceBefore=15, spaceAfter=8),
            2: ParagraphStyle("ReportH2", parent=normal, fontSize=14, leading=21, textColor=colors.HexColor("#17233A"), spaceBefore=13, spaceAfter=7),
            3: ParagraphStyle("ReportH3", parent=normal, fontSize=12, leading=19, textColor=colors.HexColor("#25324A"), spaceBefore=11, spaceAfter=6),
            4: ParagraphStyle("ReportH4", parent=normal, fontSize=10.5, leading=17, textColor=colors.HexColor("#25324A"), spaceBefore=9, spaceAfter=5),
        }
        quote_style = ParagraphStyle("ReportQuote", parent=normal, leftIndent=10, rightIndent=6, borderColor=colors.HexColor("#5B7CFA"), borderWidth=1, borderPadding=8, backColor=colors.HexColor("#F5F7FB"))
        code_style = ParagraphStyle("ReportCode", parent=normal, fontName=font_name, fontSize=8.5, leading=14, leftIndent=8, rightIndent=8, borderPadding=8, backColor=colors.HexColor("#F1F4F9"))
        story: list[Any] = [
            Paragraph(self._inline_pdf(title), title_style),
            Paragraph("生成时间：" + datetime.now().strftime("%Y-%m-%d %H:%M"), meta_style),
            HRFlowable(width="100%", thickness=1, color=colors.HexColor("#DDE3EC"), spaceAfter=12),
        ]
        for block in blocks:
            if block.kind == "heading":
                story.append(Paragraph(self._inline_pdf(block.text), heading_styles[min(max(block.level, 1), 4)]))
            elif block.kind == "paragraph":
                story.append(Paragraph(self._inline_pdf(block.text), normal))
            elif block.kind == "quote":
                story.append(Paragraph(self._inline_pdf(block.text), quote_style))
            elif block.kind == "rule":
                story.append(HRFlowable(width="100%", thickness=0.7, color=colors.HexColor("#DDE3EC"), spaceBefore=8, spaceAfter=8))
            elif block.kind == "code":
                story.append(Paragraph(html.escape(block.text).replace("\n", "<br/>"), code_style))
            elif block.kind == "list":
                items = [ListItem(Paragraph(self._inline_pdf(item), normal), leftIndent=8) for item in block.items]
                story.append(ListFlowable(items, bulletType="1" if block.ordered else "bullet", start="1", leftIndent=18, bulletFontName=font_name, bulletFontSize=9))
                story.append(Spacer(1, 5))
            elif block.kind == "table" and block.rows:
                column_count = max(len(row) for row in block.rows)
                table_data = []
                for row in block.rows:
                    padded = row + [""] * (column_count - len(row))
                    table_data.append([Paragraph(self._inline_pdf(item), normal) for item in padded])
                available_width = A4[0] - 36 * mm
                table = Table(table_data, colWidths=[available_width / column_count] * column_count, repeatRows=1, hAlign="LEFT")
                table.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EEF2F7")),
                    ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#243044")),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#DDE3EC")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 7),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                    ("TOPPADDING", (0, 0), (-1, -1), 6),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ]))
                story.extend([table, Spacer(1, 8)])

        def footer(canvas, doc):
            canvas.saveState()
            canvas.setFont(font_name, 8)
            canvas.setFillColor(colors.HexColor("#8C94A2"))
            canvas.drawCentredString(A4[0] / 2, 9 * mm, "第 %s 页 · AI 生成报告，请结合真实业务情况复核" % doc.page)
            canvas.restoreState()

        document.build(story, onFirstPage=footer, onLaterPages=footer)
        return output.getvalue()

    def _add_docx_inline(self, paragraph, value: str, bold: bool = False) -> None:
        pattern = re.compile(r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\]\([^)]+\))")
        position = 0
        for match in pattern.finditer(str(value or "")):
            if match.start() > position:
                run = paragraph.add_run(value[position:match.start()])
                run.bold = bold
            token = match.group(0)
            if token.startswith("**"):
                run = paragraph.add_run(token[2:-2])
                run.bold = True
            elif token.startswith("`"):
                run = paragraph.add_run(token[1:-1])
                run.font.name = "Consolas"
            else:
                link = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
                text = "%s（%s）" % (link.group(1), link.group(2)) if link else token
                run = paragraph.add_run(text)
                run.bold = bold
            position = match.end()
        if position < len(value):
            run = paragraph.add_run(value[position:])
            run.bold = bold

    def _inline_html(self, value: str) -> str:
        text = html.escape(str(value or ""))
        text = re.sub(r"\[([^\]]+)\]\((https?://[^\s)]+)\)", r'<a href="\2">\1</a>', text)
        text = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text)
        text = re.sub(r"`(.+?)`", r"<code>\1</code>", text)
        return text

    def _inline_pdf(self, value: str) -> str:
        text = html.escape(str(value or ""))
        text = re.sub(r"\[([^\]]+)\]\((https?://[^\s)]+)\)", r'<link href="\2" color="#315DCC">\1</link>', text)
        text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
        text = re.sub(r"`(.+?)`", r"<font backColor=\"#F1F4F9\">\1</font>", text)
        return text


report_renderer = ReportRenderer()
